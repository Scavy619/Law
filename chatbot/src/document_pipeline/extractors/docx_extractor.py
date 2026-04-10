import io
from docx import Document


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a .docx file.

    Args:
        file_bytes: Raw bytes of the uploaded .docx file

    Returns:
        Extracted text as string

    Raises:
        ValueError if file is empty or unreadable
    """
    doc = Document(io.BytesIO(file_bytes))

    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    # Tables ka text bhi extract karo (legal docs mein tables common hain)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)

    if not paragraphs:
        raise ValueError("DOCX file is either empty or unreadable.")

    return "\n\n".join(paragraphs)